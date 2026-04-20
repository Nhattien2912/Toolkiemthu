from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = PROJECT_ROOT / "Nhom03_FinalProject.docx"
LEGACY_TEMPLATE = PROJECT_ROOT / "Nhom03_NgoQuangLoi_TrinhNhatTien_NguyenCongQuoc.docx"


def _add_kv_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(value)


def generate_word_report(
    output_path: Path,
    latest_run: dict | None,
    cumulative_coverage: dict | None,
    bugs: list[dict],
    project_doc: dict,
    template_path: Path | None = None,
) -> Path:
    template = template_path or DEFAULT_TEMPLATE
    if not template.exists() and template == DEFAULT_TEMPLATE:
        template = LEGACY_TEMPLATE
    if template.exists():
        document = Document(str(template))
    else:
        document = Document()

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("PHỤ LỤC TỰ ĐỘNG TỪ DASHBOARD TESTOPS PRO", level=1)
    document.add_paragraph(
        "Phần này được sinh tự động từ dashboard để bổ sung số liệu thực tế cho báo cáo Word."
    )

    today_display = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    _add_kv_paragraph(document, "Thời điểm xuất báo cáo", today_display)
    _add_kv_paragraph(document, "Đề tài", project_doc.get("title", "Xây dựng kiểm thử website bán giày"))

    if latest_run:
        document.add_heading("1. Kết quả lần chạy mới nhất", level=2)
        _add_kv_paragraph(document, "XML run", latest_run["xml_filename"])
        _add_kv_paragraph(document, "Tổng test", str(latest_run["tests"]))
        _add_kv_paragraph(document, "Passed", str(latest_run["passed"]))
        _add_kv_paragraph(document, "Failed", str(latest_run["failures"] + latest_run["errors"]))
        _add_kv_paragraph(document, "Skipped", str(latest_run["skipped"]))
        _add_kv_paragraph(document, "Pass rate", f"{latest_run['pass_rate']}%")
        _add_kv_paragraph(document, "Thời gian chạy", latest_run["duration_display"])

    if cumulative_coverage:
        document.add_heading("2. Độ phủ tích lũy", level=2)
        _add_kv_paragraph(document, "Tổng test case Excel", str(cumulative_coverage["total_cases"]))
        _add_kv_paragraph(document, "Đã có kết quả thực", str(cumulative_coverage["executed_cases"]))
        _add_kv_paragraph(document, "Chưa thực hiện", str(cumulative_coverage["not_executed_cases"]))
        _add_kv_paragraph(document, "Passed", str(cumulative_coverage["passed_cases"]))
        _add_kv_paragraph(document, "Failed", str(cumulative_coverage["failed_cases"]))
        _add_kv_paragraph(document, "Coverage", f"{cumulative_coverage['coverage_pct']}%")

    document.add_heading("3. Danh sách bug nổi bật", level=2)
    if bugs:
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Case ID"
        header_cells[1].text = "Test name"
        header_cells[2].text = "Severity"
        header_cells[3].text = "Occurrences"
        header_cells[4].text = "Latest error"

        for bug in bugs[:20]:
            row_cells = table.add_row().cells
            row_cells[0].text = bug["case_id"]
            row_cells[1].text = bug["name"]
            row_cells[2].text = bug["severity"]
            row_cells[3].text = str(bug["occurrences"])
            row_cells[4].text = (bug["latest_error"] or "Không có message")[:250]
    else:
        document.add_paragraph("Không có bug nào trong dữ liệu hiện có.")

    document.add_heading("4. Ghi chú", level=2)
    for line in (
        "Nguồn số liệu: XML reports trong thư mục reports/.",
        "Bugs được tổng hợp từ toàn bộ lịch sử run hợp lệ.",
        "File này được tạo tự động để bổ sung vào báo cáo đồ án Word.",
    ):
        document.add_paragraph(f"• {line}")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
